from docx import Document

def generate_report(asset, threat, likelihood, impact):
    doc = Document()
    doc.add_heading("Risk Assessment Report", 0)
    doc.add_paragraph(f"Asset: {asset}")
    doc.add_paragraph(f"Threat: {threat}")
    doc.add_paragraph(f"Likelihood: {likelihood}/5")
    doc.add_paragraph(f"Impact: {impact}/5")
    score = int(likelihood) * int(impact)
    doc.add_paragraph(f"Risk Score: {score}/25")
    doc.save("risk_report.docx")

if __name__ == "__main__":
    asset = input("Enter asset name: ")
    threat = input("Enter threat: ")
    likelihood = input("Likelihood (1-5): ")
    impact = input("Impact (1-5): ")
    generate_report(asset, threat, likelihood, impact)